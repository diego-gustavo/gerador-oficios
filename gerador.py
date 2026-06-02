import os
import datetime
import re
import json
import sys
import shutil
from tkinter import messagebox, filedialog, colorchooser
import tkinter as tk

import customtkinter as ctk
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt
import openpyxl
from PIL import Image, ImageTk

# =============================================================================
# PYINSTALLER
# =============================================================================
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

EXE_DIR = os.path.dirname(os.path.abspath(sys.argv[0]))

# Windows taskbar icon
if sys.platform.startswith("win"):
    try:
        import ctypes
        myappid = "brtsorocaba.achados.oficios.1.0"
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    except Exception:
        pass


# =============================================================================
# CONSTANTS & CONFIGURATION
# =============================================================================

SCRIPT_DIR = resource_path("")
DEFAULT_CONFIG_FILE = resource_path("config_defaults.json")
DRAFT_FILE = os.path.join(EXE_DIR, "rascunho.json")
CONFIG_BACKUP_FILE = os.path.join(EXE_DIR, "config_backup.json")
CONFIG_FILE = os.path.join(EXE_DIR, "config.json")

MONTHS = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
]

DEFAULT_SUGGESTIONS = [
    "Vale Transporte", "RG", "CIN", "Carteira", "Cartão", "Celular",
    "Chave", "Documento", "Bolsas", "Mochila", "Óculos", "Fone de Ouvido",
    "Carregador", "Guarda-chuva/sol", "Jaqueta", "Outros",
]

APPEARANCE_COLOR_FIELDS = [
    ("Cor Primária (Destaques):", "color_highlight_light", "color_highlight_dark"),
    ("Cor Secundária (Botões):", "color_confirm_light", "color_confirm_dark"),
    ("Texto dos Botões Principais:", "color_confirm_text_light", "color_confirm_text_dark"),
    ("Cor Terciária (Bordas):", "color_outline_light", "color_outline_dark"),
    ("Cor do Texto Padrão:", "color_text_light", "color_text_dark"),
    ("Fundo do Aplicativo:", "color_app_bg_light", "color_app_bg_dark"),
    ("Painéis e Listas:", "color_panel_light", "color_panel_dark"),
    ("Campos de Entrada:", "color_input_bg_light", "color_input_bg_dark"),
    ("Ações de Exclusão:", "color_danger_light", "color_danger_dark"),
]

APPEARANCE_NUMBER_FIELDS = [
    ("Arredondamento dos Componentes:", "corner_radius", ["0", "4", "8", "12", "16"], "8"),
    ("Espessura das Bordas:", "border_width", ["0", "1", "2", "3"], "1"),
    ("Altura do Logo:", "logo_height", ["32", "40", "48", "56"], "40"),
]

RESTART_CONFIG_KEYS = [
    "tema",
    "escala_widget",
    "escala_texto",
    "fonte_dislexia",
    "alto_contraste",
    "fonte_familia",
    "fonte_tamanho_base",
    "fonte_negrito_padrao",
    "corner_radius",
    "border_width",
    "logo_height",
]
for _, light_key, dark_key in APPEARANCE_COLOR_FIELDS:
    RESTART_CONFIG_KEYS.extend([light_key, dark_key])


def load_config():
    """Load configuration from JSON file, merging with built-in defaults."""
    default_config = {
        "excel_path": r"J:/Usuários/Documentos Gerais/Controle Ofícios - BRT Operação.xlsx",
        "template_path": os.path.join(SCRIPT_DIR, "M:/Remissão/Base/Relatórios Fechados/Ofícios e protocolos/Ofícios 2026 - BRT/Achados e Perdidos/template.docx"),
        "save_dir": r"M:/Remissão/Base/Relatórios Fechados/Ofícios e protocolos/Ofícios 2026 - BRT",
        "tema": "Light",
        "escala": "100%",
        "fonte_dislexia": False,
        "alto_contraste": False,
        "fonte_familia": "Roboto",
        "fonte_tamanho_base": 12,
        "fonte_negrito_padrao": True,
        "escala_widget": "100%",
        "escala_texto": "100%",
        "sugestoes_itens": DEFAULT_SUGGESTIONS,
        "color_highlight_light": "#1a1264",
        "color_highlight_dark": "#7b85f0",
        "color_confirm_light": "#39b44a",
        "color_confirm_dark": "#2e933b",
        "color_outline_light": "#6d6e71",
        "color_outline_dark": "#8b8c8f",
        "color_text_light": "#000000",
        "color_text_dark": "#ffffff",
        "color_confirm_text_light": "#ffffff",
        "color_confirm_text_dark": "#ffffff",
        "color_app_bg_light": "#f7f7f7",
        "color_app_bg_dark": "#1f1f1f",
        "color_panel_light": "#ffffff",
        "color_panel_dark": "#2b2b2b",
        "color_input_bg_light": "#ffffff",
        "color_input_bg_dark": "#343638",
        "color_danger_light": "#d9534f",
        "color_danger_dark": "#ff6b6b",
        "corner_radius": "8",
        "border_width": "1",
        "logo_height": "40",
    }
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                default_config.update(json.load(f))
        except Exception:
            pass
    return default_config


def save_config(config):
    """Persist configuration dictionary to CONFIG_FILE."""
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=4)


CONFIG = load_config()
SUGGESTIONS = CONFIG["sugestoes_itens"]


def get_color_pair(light_key, dark_key, fallback_light, fallback_dark):
    """Return (light_color, dark_color) tuple from current configuration."""
    return (
        CONFIG.get(light_key, fallback_light),
        CONFIG.get(dark_key, fallback_dark),
    )


def get_config_int(key, fallback):
    """Safely retrieve an integer configuration value."""
    try:
        return int(CONFIG.get(key, fallback))
    except (TypeError, ValueError):
        return int(fallback)


COLOR_HIGHLIGHT = get_color_pair("color_highlight_light", "color_highlight_dark", "#1a1264", "#7b85f0")
COLOR_CONFIRM = get_color_pair("color_confirm_light", "color_confirm_dark", "#39b44a", "#2e933b")
COLOR_CONFIRM_TEXT = get_color_pair("color_confirm_text_light", "color_confirm_text_dark", "#ffffff", "#ffffff")
COLOR_OUTLINE = get_color_pair("color_outline_light", "color_outline_dark", "#6d6e71", "#8b8c8f")
COLOR_TEXT = get_color_pair("color_text_light", "color_text_dark", "#000000", "#ffffff")
COLOR_APP_BG = get_color_pair("color_app_bg_light", "color_app_bg_dark", "#f7f7f7", "#1f1f1f")
COLOR_PANEL = get_color_pair("color_panel_light", "color_panel_dark", "#ffffff", "#2b2b2b")
COLOR_INPUT_BG = get_color_pair("color_input_bg_light", "color_input_bg_dark", "#ffffff", "#343638")
COLOR_DANGER = get_color_pair("color_danger_light", "color_danger_dark", "#d9534f", "#ff6b6b")
BORDER_WIDTH = get_config_int("border_width", 1)
CORNER_RADIUS = get_config_int("corner_radius", 8)
LOGO_HEIGHT = get_config_int("logo_height", 40)

ctk.set_appearance_mode(CONFIG["tema"])
try:
    widget_scale = int(CONFIG.get("escala_widget", "100%").replace("%", "")) / 100.0
    ctk.set_widget_scaling(widget_scale)
except:
    pass

# High contrast mode overrides
if CONFIG.get("alto_contraste"):
    COLOR_HIGHLIGHT = ("#000000", "#FFFFFF")
    COLOR_CONFIRM = "transparent"
    COLOR_CONFIRM_TEXT = ("#000000", "#FFFFFF")
    COLOR_OUTLINE = ("#000000", "#FFFFFF")
    COLOR_TEXT = ("#000000", "#FFFFFF")
    COLOR_APP_BG = ("#FFFFFF", "#000000")
    COLOR_PANEL = ("#FFFFFF", "#000000")
    COLOR_INPUT_BG = ("#FFFFFF", "#000000")
    COLOR_DANGER = ("#000000", "#FFFFFF")
    BORDER_WIDTH = 2

    for key in ["CTk", "CTkFrame", "CTkScrollableFrame", "CTkTabview", "CTkLabel", "CTkSegmentedButton"]:
        if key in ctk.ThemeManager.theme:
            if "fg_color" in ctk.ThemeManager.theme[key]:
                ctk.ThemeManager.theme[key]["fg_color"] = ("#FFFFFF", "#000000")
            if "text_color" in ctk.ThemeManager.theme[key]:
                ctk.ThemeManager.theme[key]["text_color"] = ("#000000", "#FFFFFF")
    if "CTkScrollableFrame" in ctk.ThemeManager.theme:
        ctk.ThemeManager.theme["CTkScrollableFrame"]["label_fg_color"] = ("#FFFFFF", "#000000")


def apply_custom_theme():
    """Overwrite CustomTkinter's default theme with the current palette."""
    adjustments = {
        "CTk": {"fg_color": COLOR_APP_BG},
        "CTkFrame": {
            "fg_color": COLOR_PANEL,
            "border_color": COLOR_OUTLINE,
            "corner_radius": CORNER_RADIUS,
            "border_width": BORDER_WIDTH,
        },
        "CTkScrollableFrame": {
            "fg_color": COLOR_PANEL,
            "border_color": COLOR_OUTLINE,
            "corner_radius": CORNER_RADIUS,
            "border_width": BORDER_WIDTH,
        },
        "CTkEntry": {
            "fg_color": COLOR_INPUT_BG,
            "text_color": COLOR_TEXT,
            "border_color": COLOR_OUTLINE,
            "corner_radius": CORNER_RADIUS,
            "border_width": BORDER_WIDTH,
        },
        "CTkComboBox": {
            "fg_color": COLOR_INPUT_BG,
            "text_color": COLOR_TEXT,
            "border_color": COLOR_OUTLINE,
            "button_color": COLOR_CONFIRM,
            "button_hover_color": COLOR_HIGHLIGHT,
            "corner_radius": CORNER_RADIUS,
            "border_width": BORDER_WIDTH,
        },
        "CTkButton": {
            "border_color": COLOR_OUTLINE,
            "corner_radius": CORNER_RADIUS,
            "border_width": BORDER_WIDTH,
        },
        "CTkLabel": {"text_color": COLOR_TEXT},
        "CTkSwitch": {
            "text_color": COLOR_TEXT,
            "progress_color": COLOR_CONFIRM,
            "button_color": COLOR_HIGHLIGHT,
            "button_hover_color": COLOR_HIGHLIGHT,
        },
        "CTkTabview": {
            "fg_color": COLOR_PANEL,
            "border_color": COLOR_OUTLINE,
            "corner_radius": CORNER_RADIUS,
            "border_width": BORDER_WIDTH,
        },
    }
    for widget, options in adjustments.items():
        if widget in ctk.ThemeManager.theme:
            ctk.ThemeManager.theme[widget].update(options)


apply_custom_theme()


def get_font(size=None, bold=False):
    """Build a font tuple respecting accessibility settings."""
    base_size = CONFIG.get("fonte_tamanho_base", 12)
    if size is None:
        size = base_size
    try:
        text_scale = int(CONFIG.get("escala_texto", "100%").replace("%", "")) / 100.0
    except:
        text_scale = 1.0
    scaled_size = int(size * text_scale)
    family = "Comic Sans MS" if CONFIG.get("fonte_dislexia") else CONFIG.get("fonte_familia", "Roboto")
    if CONFIG.get("fonte_negrito_padrao") and not bold:
        bold = True
    weight = "bold" if bold else "normal"
    return (family, scaled_size, weight)


# =============================================================================
# WORD / EXCEL HELPERS
# =============================================================================

def date_to_long_text(date=None):
    """Convert a date object into Brazilian long format (e.g., '1 de Janeiro de 2026')."""
    if date is None:
        date = datetime.date.today()
    return f"{date.day} de {MONTHS[date.month - 1]} de {date.year}"


def find_sheet_by_year(workbook, year):
    """Return the sheet name that contains '<year> Ofícios Emitidos'."""
    pattern = re.compile(rf"{year}.*Of[ií]cios\s+Emitidos", re.IGNORECASE)
    for name in workbook.sheetnames:
        if pattern.search(name):
            return name
    raise ValueError(
        f"Nenhuma aba com ano {year} e 'Ofícios Emitidos' encontrada.\nAbas disponíveis: {workbook.sheetnames}"
    )


def get_next_number(worksheet, year_ref):
    """Calculate the next sequential officio number for the given year."""
    largest = 0
    for row in worksheet.iter_rows(min_row=1, max_col=1, values_only=True):
        cell_val = row[0]
        if cell_val and isinstance(cell_val, str):
            match = re.match(r"(\d+)/(\d{4})", cell_val.strip())
            if match:
                num = int(match.group(1))
                cell_year = int(match.group(2))
                if cell_year == year_ref and num > largest:
                    largest = num
    next_num = largest + 1
    return f"{next_num}/{year_ref}"


def find_last_row(worksheet, column=1):
    """Return the last non-empty row index in the given column."""
    max_row = worksheet.max_row
    while max_row > 0 and worksheet.cell(row=max_row, column=column).value is None:
        max_row -= 1
    return max_row


def ensure_list_style(doc):
    """Create a 'List Number' style in the document if it does not exist."""
    try:
        doc.styles["List Number"]
    except KeyError:
        from lxml import etree

        numbering_part = doc.part.numbering_part
        numbering_xml = numbering_part.numbering_definitions._numbering

        abstract_num_xml = """
        <w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="10">
            <w:multiLevelType w:val="hybridMultilevel"/>
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="720" w:hanging="360"/>
                </w:pPr>
            </w:lvl>
        </w:abstractNum>
        """
        num_xml = """
        <w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="10">
            <w:abstractNumId w:val="10"/>
        </w:num>
        """
        numbering_xml.append(etree.fromstring(abstract_num_xml))
        numbering_xml.append(etree.fromstring(num_xml))

        style_xml = """
        <w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="paragraph" w:styleId="ListNumber">
            <w:name w:val="List Number"/>
            <w:basedOn w:val="Normal"/>
            <w:pPr>
                <w:numPr>
                    <w:ilvl w:val="0"/>
                    <w:numId w:val="10"/>
                </w:numPr>
            </w:pPr>
        </w:style>
        """
        doc.styles.element.append(etree.fromstring(style_xml))


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph right after the given one."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    if style:
        new_para.style = style
    return new_para


def replace_tag_in_doc(document, tag, new_text):
    """Replace all occurrences of `tag` with `new_text` in the document."""
    for paragraph in document.paragraphs:
        if tag in paragraph.text:
            for run in paragraph.runs:
                if tag in run.text:
                    run.text = run.text.replace(tag, new_text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if tag in paragraph.text:
                        for run in paragraph.runs:
                            if tag in run.text:
                                run.text = run.text.replace(tag, new_text)


def generate_document(officio_number, items, officio_date, save_path):
    """Fill the Word template and save the final document."""
    doc = Document(CONFIG["template_path"])
    long_date = date_to_long_text(officio_date)
    replace_tag_in_doc(doc, "{{DATA}}", long_date)
    replace_tag_in_doc(doc, "{{OFICIO}}", officio_number)

    # Locate the list placeholder
    list_paragraph = None
    for p in doc.paragraphs:
        if "{{LISTA_ITENS}}" in p.text:
            list_paragraph = p
            break
    if list_paragraph is None:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{{LISTA_ITENS}}" in p.text:
                            list_paragraph = p
                            break
        if list_paragraph is None:
            raise ValueError("Tag {{LISTA_ITENS}} não encontrada no template.")

    font_run = list_paragraph.runs[0].font if list_paragraph.runs else None
    ensure_list_style(doc)

    # Prepare the list paragraph
    list_paragraph.text = ""
    list_paragraph.style = doc.styles["List Number"]
    pPr = list_paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "10")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

    prev_paragraph = list_paragraph
    for idx, item in enumerate(items):
        name, code, obs = item["nome"], item["codigo"], item["obs"]
        line = name
        if code:
            line += f" – {code}"
        if obs:
            line += f" ({obs})"

        if idx == 0:
            run = list_paragraph.add_run(line)
            para = list_paragraph
        else:
            new_p = insert_paragraph_after(prev_paragraph, style="List Number")
            pPr = new_p._element.get_or_add_pPr()
            numPr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "10")
            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)
            run = new_p.add_run(line)
            prev_paragraph = new_p
            para = new_p

        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0

        if font_run:
            if font_run.name:
                run.font.name = font_run.name
            if font_run.size:
                run.font.size = font_run.size
            if font_run.bold:
                run.font.bold = True
            if font_run.italic:
                run.font.italic = True
            if font_run.color and font_run.color.rgb:
                run.font.color.rgb = font_run.color.rgb

    doc.save(save_path)
    return save_path


def add_excel_row(officio_number, year, officio_date, responsible):
    """Append a new row to the Excel control workbook."""
    wb = openpyxl.load_workbook(CONFIG["excel_path"])
    sheet_name = find_sheet_by_year(wb, year)
    ws = wb[sheet_name]

    last_row = find_last_row(ws, column=1)
    new_row = last_row + 1 if last_row > 0 else 1

    ws.cell(row=new_row, column=1, value=officio_number)
    ws.cell(row=new_row, column=2, value="Encaminhamento de Achados e Perdidos")
    ws.cell(row=new_row, column=3, value=officio_date.strftime("%d/%m/%Y"))
    ws.cell(row=new_row, column=4, value="Urbes")
    ws.cell(row=new_row, column=5, value=responsible)

    try:
        wb.save(CONFIG["excel_path"])
    except PermissionError:
        messagebox.showerror(
            "Erro",
            "Planilha de controle está aberta em outro programa. Feche-a e tente novamente.",
        )
        return False
    return True


# =============================================================================
# MAIN APPLICATION WINDOW
# =============================================================================

class OficioGeneratorApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Ofícios Achados e Perdidos - BRT Sorocaba")
        self.geometry("850x700")
        self.minsize(600, 500)
        self.configure(fg_color=COLOR_APP_BG)

        self.items = []
        self.nextOfficio = ctk.StringVar(value="Carregando...")
        self.selectedYear = ctk.StringVar(value=str(datetime.date.today().year))
        self.textHlColor = COLOR_HIGHLIGHT

        # ---- App icon ----
        icon_path = os.path.join(SCRIPT_DIR, "img", "favicon.png")
        icon_ico_path = os.path.join(SCRIPT_DIR, "img", "favicon.ico")
        if os.path.exists(icon_path):
            try:
                if not os.path.exists(icon_ico_path):
                    img_pil = Image.open(icon_path)
                    img_pil.save(icon_ico_path, format="ICO", sizes=[(32, 32)])
                self.iconbitmap(icon_ico_path)
            except Exception as e:
                print("Erro ao aplicar favicon (.ico):", e)
            try:
                self._icon_img = ImageTk.PhotoImage(Image.open(icon_path))
                self.iconphoto(True, self._icon_img)
            except Exception as e:
                print("Erro ao aplicar favicon interno:", e)

        # ---- Header (no borders) ----
        self.headerFrame = ctk.CTkFrame(self, fg_color="transparent", border_width=0)
        self.headerFrame.pack(fill="x", padx=20, pady=(15, 0))

        self.lblTitle = ctk.CTkLabel(
            self.headerFrame,
            text="Gerador de Ofícios - BRT",
            font=get_font(20, True),
            text_color=self.textHlColor,
        )
        self.lblTitle.pack(side="left", anchor="w")

        logo_path = os.path.join(SCRIPT_DIR, "img", "logo.png")
        logo_dark_path = os.path.join(SCRIPT_DIR, "img", "logo-branco.png")
        if os.path.exists(logo_path):
            try:
                pil_logo = Image.open(logo_path)
                pil_logo_dark = Image.open(logo_dark_path) if os.path.exists(logo_dark_path) else pil_logo
                ratio = LOGO_HEIGHT / float(pil_logo.size[1])
                new_width = int(float(pil_logo.size[0]) * ratio)
                logo_img = ctk.CTkImage(
                    light_image=pil_logo,
                    dark_image=pil_logo_dark,
                    size=(new_width, LOGO_HEIGHT),
                )
                self.lblLogo = ctk.CTkLabel(self.headerFrame, image=logo_img, text="")
                self.lblLogo.pack(side="right", anchor="e")
            except Exception:
                pass

        # ---- Tabs ----
        self.tabview = ctk.CTkTabview(
            self,
            fg_color=COLOR_PANEL,
            border_color=COLOR_OUTLINE,
            border_width=BORDER_WIDTH,
            corner_radius=CORNER_RADIUS,
        )
        self.tabview.pack(fill="both", expand=True, padx=20, pady=20)

        self.generatorTab = self.tabview.add("Gerador")
        self.configTab = self.tabview.add("Configurações")
        self.helpTab = self.tabview.add("Ajuda")

        self.setup_generator_tab()
        self.setup_configuration_tab()
        self.setup_help_tab()

        self.load_next_number()
        self.after(500, self._check_config_backup)
        self.after(600, self.auto_check_draft)

    # -------------------------------------------------------------------------
    # UI utility methods
    # -------------------------------------------------------------------------
    def bind_enter_key(self, widget):
        widget.bind("<Return>", lambda e: widget.tk_focusNext().focus_set())

    def create_section_header(self, parent, title, pady=(20, 5), font_size=16):
        ctk.CTkLabel(
            parent,
            text=title,
            font=get_font(font_size, True),
            text_color=self.textHlColor,
        ).pack(pady=pady, anchor="w")

    def create_config_row(self, parent, label_text, label_width=160, pady=5):
        frame = ctk.CTkFrame(parent, fg_color="transparent", border_width=0)
        frame.pack(fill="x", pady=pady)
        ctk.CTkLabel(frame, text=label_text, width=label_width, anchor="w", text_color=COLOR_TEXT).pack(side="left")
        return frame

    def create_config_combobox(self, parent, label_text, values, current, label_width=160):
        frame = self.create_config_row(parent, label_text, label_width=label_width)
        combo = ctk.CTkComboBox(
            frame,
            values=values,
            text_color=COLOR_TEXT,
            fg_color=COLOR_INPUT_BG,
            border_color=COLOR_OUTLINE,
            corner_radius=CORNER_RADIUS,
        )
        combo.pack(side="left")
        combo.set(current)
        return combo

    def create_config_path(self, parent, label_text, value, browse_command):
        frame = self.create_config_row(parent, label_text, label_width=120)
        entry = ctk.CTkEntry(frame, text_color=COLOR_TEXT, fg_color=COLOR_INPUT_BG)
        entry.pack(side="left", fill="x", expand=True, padx=5)
        entry.insert(0, value)
        ctk.CTkButton(frame, text="Procurar", width=80, command=lambda: browse_command(entry)).pack(side="left")
        return entry

    def create_config_switch(self, parent, text, selected=False):
        frame = ctk.CTkFrame(parent, fg_color="transparent", border_width=0)
        frame.pack(fill="x", pady=5)
        switch = ctk.CTkSwitch(frame, text=text, text_color=COLOR_TEXT)
        switch.pack(side="left")
        if selected:
            switch.select()
        return switch

    # -------------------------------------------------------------------------
    # Generator tab
    # -------------------------------------------------------------------------
    def setup_generator_tab(self):
        scroll = ctk.CTkScrollableFrame(self.generatorTab, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=5, pady=5)

        # Top frame: year + next officio
        topFrame = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        topFrame.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(topFrame, text="Ano:", font=get_font(14, True), text_color=self.textHlColor).pack(side="left", padx=5)

        # Year control with up/down buttons
        self.yearFrame = ctk.CTkFrame(topFrame, fg_color="transparent", border_width=0)
        self.yearFrame.pack(side="left", padx=5)

        self.btnYearDown = ctk.CTkButton(
            self.yearFrame, text="◀", width=25, height=25,
            command=self.decrease_year,
            fg_color="transparent", border_width=0,
            text_color=COLOR_TEXT, font=get_font(10)
        )
        self.btnYearDown.pack(side="left")

        self.lblYear = ctk.CTkLabel(
            self.yearFrame, textvariable=self.selectedYear,
            width=50, font=get_font(14, True),
            fg_color="transparent", text_color=COLOR_TEXT
        )
        self.lblYear.pack(side="left")

        self.btnYearUp = ctk.CTkButton(
            self.yearFrame, text="▶", width=25, height=25,
            command=self.increase_year,
            fg_color="transparent", border_width=0,
            text_color=COLOR_TEXT, font=get_font(10)
        )
        self.btnYearUp.pack(side="left")

        ctk.CTkLabel(topFrame, text="Próximo Ofício:", font=get_font(14, True), text_color=self.textHlColor).pack(side="left", padx=(20, 5))
        self.lblNextOfficio = ctk.CTkLabel(topFrame, textvariable=self.nextOfficio, font=get_font(14, True))
        self.lblNextOfficio.pack(side="left", padx=5)

        sep = ctk.CTkFrame(scroll, height=2, fg_color=COLOR_OUTLINE, border_width=0)
        sep.pack(fill="x", pady=10)

        # Form
        formFrame = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        formFrame.pack(fill="x")

        ctk.CTkLabel(formFrame, text="Item:", font=get_font(12), text_color=COLOR_TEXT).grid(row=0, column=0, sticky="w", pady=5, padx=5)
        self.itemCombobox = ctk.CTkComboBox(formFrame, values=SUGGESTIONS, width=300, border_color=COLOR_OUTLINE, text_color=COLOR_TEXT)
        self.itemCombobox.grid(row=0, column=1, sticky="w", pady=5, padx=5)
        self.itemCombobox.set("")
        self.bind_enter_key(self.itemCombobox._entry)

        ctk.CTkLabel(formFrame, text="ID/Código:", font=get_font(12), text_color=COLOR_TEXT).grid(row=1, column=0, sticky="w", pady=5, padx=5)
        self.codeEntry = ctk.CTkEntry(formFrame, width=300, border_color=COLOR_OUTLINE, text_color=COLOR_TEXT)
        self.codeEntry.grid(row=1, column=1, sticky="w", pady=5, padx=5)
        self.bind_enter_key(self.codeEntry)

        ctk.CTkLabel(formFrame, text="Cor/Obs:", font=get_font(12), text_color=COLOR_TEXT).grid(row=2, column=0, sticky="w", pady=5, padx=5)
        self.obsEntry = ctk.CTkEntry(formFrame, width=300, border_color=COLOR_OUTLINE, text_color=COLOR_TEXT)
        self.obsEntry.grid(row=2, column=1, sticky="w", pady=5, padx=5)
        self.bind_enter_key(self.obsEntry)

        ctk.CTkLabel(formFrame, text="Data do Ofício:", font=get_font(12), text_color=COLOR_TEXT).grid(row=3, column=0, sticky="w", pady=5, padx=5)
        self.dateEntry = ctk.CTkEntry(formFrame, width=300, border_color=COLOR_OUTLINE, text_color=COLOR_TEXT)
        self.dateEntry.insert(0, datetime.date.today().strftime("%d/%m/%Y"))
        self.dateEntry.grid(row=3, column=1, sticky="w", pady=5, padx=5)
        self.bind_enter_key(self.dateEntry)

        ctk.CTkLabel(formFrame, text="Responsável:", font=get_font(12), text_color=COLOR_TEXT).grid(row=4, column=0, sticky="w", pady=5, padx=5)
        self.respEntry = ctk.CTkEntry(formFrame, width=300, border_color=COLOR_OUTLINE, text_color=COLOR_TEXT)
        self.respEntry.grid(row=4, column=1, sticky="w", pady=5, padx=5)
        self.respEntry.bind("<Return>", lambda e: self.add_item())

        self.addButton = ctk.CTkButton(
            formFrame,
            text="Adicionar Item",
            command=self.add_item,
            font=get_font(12, True),
            border_width=0,
            fg_color="transparent",
            text_color=COLOR_TEXT,
        )
        self.addButton.grid(row=5, column=0, columnspan=2, pady=15)

        # Items list
        listFrame = ctk.CTkFrame(scroll, fg_color=COLOR_PANEL, corner_radius=CORNER_RADIUS,
                                 border_width=BORDER_WIDTH, border_color=COLOR_OUTLINE)
        listFrame.pack(fill="both", expand=True, pady=5)

        ctk.CTkLabel(listFrame, text="Itens Adicionados:", font=get_font(14, True), text_color=self.textHlColor).pack(anchor="w", padx=10, pady=5)

        self.itemsScroll = ctk.CTkScrollableFrame(
            listFrame,
            border_width=BORDER_WIDTH,
            border_color=COLOR_OUTLINE,
            fg_color=COLOR_PANEL,
            corner_radius=CORNER_RADIUS,
            height=160,
        )
        self.itemsScroll.pack(fill="both", expand=True, padx=10, pady=5)

        # Bottom buttons (with borders)
        self.statusVar = ctk.StringVar(value="Pronto")
        ctk.CTkLabel(scroll, textvariable=self.statusVar, font=get_font(10), text_color=COLOR_TEXT).pack(anchor="w", padx=5, pady=(0, 5))

        btnFrame = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        btnFrame.pack(fill="x", pady=10)

        self.btnClear = ctk.CTkButton(
            btnFrame,
            text="Limpar Lista",
            command=self.clear_items,
            fg_color="transparent",
            border_width=BORDER_WIDTH,
            border_color=COLOR_OUTLINE,
            text_color=COLOR_TEXT,
        )
        self.btnClear.pack(side="left", padx=5)

        self.btnSaveDraft = ctk.CTkButton(
            btnFrame,
            text="Salvar Rascunho",
            command=self.save_draft,
            fg_color="transparent",
            border_width=BORDER_WIDTH,
            border_color=COLOR_OUTLINE,
            text_color=COLOR_TEXT,
        )
        self.btnSaveDraft.pack(side="left", padx=5)

        self.btnLoadDraft = ctk.CTkButton(
            btnFrame,
            text="Carregar Rascunho",
            command=self.load_draft,
            fg_color="transparent",
            border_width=BORDER_WIDTH,
            border_color=COLOR_OUTLINE,
            text_color=COLOR_TEXT,
        )
        self.btnLoadDraft.pack(side="left", padx=5)

        self.btnDeleteDraft = ctk.CTkButton(
            btnFrame,
            text="Excluir Rascunho",
            command=self.delete_draft,
            fg_color="transparent",
            border_width=BORDER_WIDTH,
            border_color=COLOR_DANGER if not CONFIG.get("alto_contraste") else COLOR_OUTLINE,
            text_color=(COLOR_DANGER if not CONFIG.get("alto_contraste") else COLOR_TEXT),
        )
        self.btnDeleteDraft.pack(side="left", padx=5)

        self.btnGenerate = ctk.CTkButton(
            btnFrame,
            text="Gerar Ofício",
            command=self.generate_document_action,
            fg_color=COLOR_CONFIRM,
            border_width=BORDER_WIDTH,
            border_color=COLOR_OUTLINE,
            text_color=COLOR_CONFIRM_TEXT,
            font=get_font(12, True),
        )
        self.btnGenerate.pack(side="right", padx=5)

    # -------------------------------------------------------------------------
    # Configuration tab (no borders on non-entry widgets)
    # -------------------------------------------------------------------------
    def setup_configuration_tab(self):
        scroll = ctk.CTkScrollableFrame(self.configTab, fg_color="transparent", border_width=0)
        scroll.pack(fill="both", expand=True, padx=5, pady=5)

        # Paths
        ctk.CTkLabel(scroll, text="Caminhos de Arquivos", font=get_font(16, True), text_color=self.textHlColor).pack(pady=(5, 10), anchor="w")

        self.entryExcel = self.create_config_path(scroll, "Planilha Excel:", CONFIG.get("excel_path", ""), self.browse_file)
        self.entryTemplate = self.create_config_path(scroll, "Template Word:", CONFIG.get("template_path", ""), self.browse_file)
        self.entrySaveDir = self.create_config_path(scroll, "Pasta Padrão:", CONFIG.get("save_dir", ""), self.browse_folder)

        # Suggestions
        ctk.CTkLabel(scroll, text="Preferências", font=get_font(16, True), text_color=self.textHlColor).pack(pady=(20, 5), anchor="w")
        f4 = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        f4.pack(fill="x", pady=5)
        ctk.CTkLabel(f4, text="Sugestões de Itens (separados por vírgula):", anchor="w").pack(side="top", anchor="w")
        self.entrySuggestions = ctk.CTkEntry(f4)
        self.entrySuggestions.pack(side="top", fill="x", expand=True)
        self.entrySuggestions.insert(0, ", ".join(CONFIG.get("sugestoes_itens", [])))

        # Accessibility
        ctk.CTkLabel(scroll, text="Acessibilidade", font=get_font(16, True), text_color=self.textHlColor).pack(pady=(20, 5), anchor="w")

        self.comboTheme = self.create_config_combobox(scroll, "Tema Visual:", ["Light", "Dark", "System"], CONFIG.get("tema", "System"))
        self.comboWidgetScale = self.create_config_combobox(scroll, "Tamanho da Interface:", ["80%", "100%", "120%", "150%"], CONFIG.get("escala_widget", "100%"))
        self.comboTextScale = self.create_config_combobox(scroll, "Tamanho do Texto:", ["80%", "100%", "120%", "150%", "200%"], CONFIG.get("escala_texto", "100%"))

        self.switchDyslexia = self.create_config_switch(scroll, "Usar fonte para Dislexia (Comic Sans MS)", CONFIG.get("fonte_dislexia"))
        self.switchHighContrast = self.create_config_switch(scroll, "Modo Alto Contraste", CONFIG.get("alto_contraste"))

        ctk.CTkLabel(
            scroll,
            text="* Alterações de tema/escala/contraste requerem reiniciar o aplicativo.",
            font=get_font(11),
            text_color=COLOR_OUTLINE,
        ).pack(pady=(10, 5), anchor="w")

        self.setup_appearance_tab(scroll)

        # Buttons
        fbtn = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        fbtn.pack(pady=15)

        ctk.CTkButton(fbtn, text="Salvar Configurações", command=self.save_config_from_ui, fg_color=COLOR_CONFIRM, border_width=BORDER_WIDTH, border_color=COLOR_OUTLINE, text_color=COLOR_CONFIRM_TEXT).pack(side="left", padx=10)
        ctk.CTkButton(fbtn, text="Restaurar Padrões", command=self.restore_defaults, fg_color="transparent", border_width=BORDER_WIDTH, border_color=COLOR_DANGER if not CONFIG.get("alto_contraste") else COLOR_OUTLINE, text_color=(COLOR_DANGER if not CONFIG.get("alto_contraste") else COLOR_TEXT)).pack(side="left", padx=10)
        ctk.CTkButton(fbtn, text="Definir como Padrão do App", command=self.set_as_default, fg_color="transparent", border_width=BORDER_WIDTH, border_color=COLOR_OUTLINE, text_color=COLOR_TEXT).pack(side="left", padx=10)

    # -------------------------------------------------------------------------
    # Appearance subtab (inside config)
    # -------------------------------------------------------------------------
    def setup_appearance_tab(self, scroll):
        def create_color_row(parent, label_text, var_name_light, var_name_dark):
            f = ctk.CTkFrame(parent, fg_color="transparent", border_width=0)
            f.pack(fill="x", pady=5)
            ctk.CTkLabel(f, text=label_text, width=200, anchor="w", text_color=COLOR_TEXT).pack(side="left")

            val_light = CONFIG.get(var_name_light, "#ffffff")
            var_l = ctk.StringVar(value=val_light)
            setattr(self, var_name_light, var_l)

            entry_l = ctk.CTkEntry(f, width=90, textvariable=var_l, text_color=COLOR_TEXT, fg_color=COLOR_INPUT_BG, border_color=COLOR_OUTLINE, corner_radius=CORNER_RADIUS)
            entry_l.pack(side="left", padx=(5, 2))

            preview_l = ctk.CTkButton(f, text="", width=30, height=30, fg_color=val_light, hover=False, border_color=COLOR_OUTLINE, border_width=1, corner_radius=CORNER_RADIUS)
            preview_l.pack(side="left", padx=(0, 15))

            def pick_l(v=var_l, p=preview_l):
                self.open_color_picker(v, p)
                p.configure(fg_color=v.get())
            preview_l.configure(command=pick_l)

            def on_light_change(*_, v=var_l, p=preview_l):
                try:
                    val = v.get()
                    if len(val) == 7 and val.startswith("#"):
                        p.configure(fg_color=val)
                except:
                    pass
            var_l.trace_add("write", on_light_change)

            ctk.CTkLabel(f, text="Escuro:", width=50, anchor="w", text_color=COLOR_TEXT).pack(side="left")
            val_dark = CONFIG.get(var_name_dark, "#000000")
            var_d = ctk.StringVar(value=val_dark)
            setattr(self, var_name_dark, var_d)

            entry_d = ctk.CTkEntry(f, width=90, textvariable=var_d, text_color=COLOR_TEXT, fg_color=COLOR_INPUT_BG, border_color=COLOR_OUTLINE, corner_radius=CORNER_RADIUS)
            entry_d.pack(side="left", padx=(5, 2))

            preview_d = ctk.CTkButton(f, text="", width=30, height=30, fg_color=val_dark, hover=False, border_color=COLOR_OUTLINE, border_width=1, corner_radius=CORNER_RADIUS)
            preview_d.pack(side="left")

            def pick_d(v=var_d, p=preview_d):
                self.open_color_picker(v, p)
                p.configure(fg_color=v.get())
            preview_d.configure(command=pick_d)

            def on_dark_change(*_, v=var_d, p=preview_d):
                try:
                    val = v.get()
                    if len(val) == 7 and val.startswith("#"):
                        p.configure(fg_color=val)
                except:
                    pass
            var_d.trace_add("write", on_dark_change)

        ctk.CTkLabel(scroll, text="Paleta de Cores", font=get_font(16, True), text_color=self.textHlColor).pack(pady=(5, 5), anchor="w")
        ctk.CTkLabel(
            scroll,
            text="Digite o código hex ou clique no preview para abrir o seletor. O preview atualiza em tempo real.",
            font=get_font(11),
            text_color=COLOR_OUTLINE,
        ).pack(anchor="w", pady=(0, 10))

        for label_text, light_key, dark_key in APPEARANCE_COLOR_FIELDS:
            create_color_row(scroll, label_text, light_key, dark_key)

        ctk.CTkLabel(scroll, text="Componentes", font=get_font(16, True), text_color=self.textHlColor).pack(pady=(20, 5), anchor="w")

        for label_text, key, values, fallback in APPEARANCE_NUMBER_FIELDS:
            f = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
            f.pack(fill="x", pady=5)
            ctk.CTkLabel(f, text=label_text, width=220, anchor="w", text_color=COLOR_TEXT).pack(side="left")
            combo = ctk.CTkComboBox(f, values=values, width=90, text_color=COLOR_TEXT, fg_color=COLOR_INPUT_BG, border_color=COLOR_OUTLINE, corner_radius=CORNER_RADIUS)
            combo.pack(side="left")
            combo.set(str(CONFIG.get(key, fallback)))
            setattr(self, f"combo_{key}", combo)

        ctk.CTkLabel(scroll, text="Tipografia", font=get_font(16, True), text_color=self.textHlColor).pack(pady=(20, 5), anchor="w")

        ff = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        ff.pack(fill="x", pady=5)
        ctk.CTkLabel(ff, text="Fonte Principal:", width=160, anchor="w", text_color=COLOR_TEXT).pack(side="left")
        self.comboFont = ctk.CTkComboBox(ff, values=["Roboto", "Arial", "Segoe UI", "Verdana", "Tahoma", "Comic Sans MS"], text_color=COLOR_TEXT)
        self.comboFont.pack(side="left")
        self.comboFont.set(CONFIG.get("fonte_familia", "Roboto"))

        ft = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        ft.pack(fill="x", pady=5)
        ctk.CTkLabel(ft, text="Tamanho Base (px):", width=160, anchor="w", text_color=COLOR_TEXT).pack(side="left")
        self.comboFontSize = ctk.CTkComboBox(ft, values=["10", "11", "12", "13", "14", "16", "18"], text_color=COLOR_TEXT)
        self.comboFontSize.pack(side="left")
        self.comboFontSize.set(str(CONFIG.get("fonte_tamanho_base", 12)))

        fn = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
        fn.pack(fill="x", pady=5)
        self.switchBold = ctk.CTkSwitch(fn, text="Textos em negrito por padrão")
        self.switchBold.pack(side="left")
        if CONFIG.get("fonte_negrito_padrao"):
            self.switchBold.select()

        ctk.CTkLabel(
            scroll,
            text="* Alterações de aparência requerem reiniciar o aplicativo.",
            font=get_font(11),
            text_color=COLOR_OUTLINE,
        ).pack(pady=(15, 5), anchor="w")

    # -------------------------------------------------------------------------
    # Help tab (no borders at all)
    # -------------------------------------------------------------------------
    def setup_help_tab(self):
        scroll = ctk.CTkScrollableFrame(self.helpTab, fg_color="transparent", border_width=0)
        scroll.pack(fill="both", expand=True, padx=5, pady=5)

        h1 = get_font(20, bold=True)
        h2 = get_font(16, bold=True)
        bold = get_font(13, bold=True)
        normal = get_font(13)

        ctk.CTkLabel(scroll, text="Guia de Uso - Gerador de Ofícios", font=h1, text_color=self.textHlColor).pack(anchor="w", pady=(5, 10))
        ctk.CTkLabel(
            scroll,
            text="Este aplicativo facilita a criação ágil de ofícios para encaminhamento de itens de achados e perdidos.",
            font=normal,
            wraplength=700,
            justify="left",
        ).pack(anchor="w", pady=(0, 15))

        ctk.CTkLabel(scroll, text="Passo a Passo", font=h2, text_color=self.textHlColor).pack(anchor="w", pady=(5, 5))

        steps = [
            ("1. ", "Verifique o ano e o número do ofício no topo da tela."),
            ("2. ", "Preencha as informações do item encontrado (Item, ID/Código, Observação) e informe o Responsável."),
            ("3. ", "Dica: Pressione ENTER após digitar em um campo para pular automaticamente para o próximo."),
            ("4. ", "Clique em 'Adicionar Item'. Ele aparecerá na lista inferior."),
            ("5. ", "Ao adicionar todos os itens desejados, clique no botão verde 'Gerar Ofício'."),
        ]
        for num, txt in steps:
            f = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
            f.pack(fill="x", pady=2)
            ctk.CTkLabel(f, text=num, font=bold).pack(side="left", anchor="nw")
            ctk.CTkLabel(f, text=txt, font=normal, wraplength=700, justify="left").pack(side="left", anchor="nw", padx=(5, 0))

        ctk.CTkLabel(scroll, text="Funcionalidades Extras", font=h2, text_color=self.textHlColor).pack(anchor="w", pady=(15, 5))
        extras = [
            ("Rascunhos: ", "Caso precise fechar o aplicativo durante o processo, você pode salvar um Rascunho para não perder sua lista."),
            ("Configurações: ", "Na aba respectiva, você pode definir os caminhos locais para sua Planilha de Controle, o Template Word base e a pasta padrão para onde os ofícios prontos devem ir."),
            ("Acessibilidade: ", "Personalize a aparência trocando para o Modo Escuro, alterando a Escala Visual (Zoom), ativando Alto Contraste ou acionando a fonte projetada para leitura em dislexia."),
        ]
        for btext, txt in extras:
            f = ctk.CTkFrame(scroll, fg_color="transparent", border_width=0)
            f.pack(fill="x", pady=2, padx=(10, 0))
            ctk.CTkLabel(f, text="• ", font=bold).pack(side="left", anchor="nw")
            ctk.CTkLabel(f, text=btext, font=bold).pack(side="left", anchor="nw", padx=(2, 0))
            ctk.CTkLabel(f, text=txt, font=normal, wraplength=550, justify="left").pack(side="left", anchor="nw", padx=(5, 0))

    # -------------------------------------------------------------------------
    # File browsing helpers
    # -------------------------------------------------------------------------
    def browse_file(self, entry_widget):
        path = filedialog.askopenfilename()
        if path:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, path)

    def browse_folder(self, entry_widget):
        path = filedialog.askdirectory()
        if path:
            entry_widget.delete(0, "end")
            entry_widget.insert(0, path)

    # -------------------------------------------------------------------------
    # Config backup restore logic
    # -------------------------------------------------------------------------
    def _check_config_backup(self):
        if os.path.exists(CONFIG_BACKUP_FILE):
            self.lift()
            self.focus_force()
            result = messagebox.askyesno(
                "Manter Alterações?",
                "O aplicativo foi reiniciado com novas configurações visuais.\n\nDeseja manter as alterações?",
                parent=self,
            )
            if result:
                os.remove(CONFIG_BACKUP_FILE)
                self.statusVar.set("Novas configurações mantidas.")
            else:
                try:
                    with open(CONFIG_BACKUP_FILE, "r", encoding="utf-8") as f:
                        old_config = json.load(f)
                    save_config(old_config)
                    os.remove(CONFIG_BACKUP_FILE)
                    messagebox.showinfo("Revertido", "Configurações restauradas. O aplicativo será reiniciado.", parent=self)
                    self.destroy()
                    os.execl(sys.executable, sys.executable, *sys.argv)
                except Exception as e:
                    messagebox.showerror("Erro", f"Não foi possível reverter:\n{e}", parent=self)

    def restore_defaults(self):
        if os.path.exists(DEFAULT_CONFIG_FILE):
            msg = "Deseja restaurar para o padrão que você definiu anteriormente?"
        else:
            msg = "Deseja restaurar todas as configurações para o padrão de fábrica?"
        if messagebox.askyesno("Restaurar Padrões", msg):
            if os.path.exists(DEFAULT_CONFIG_FILE):
                shutil.copy2(DEFAULT_CONFIG_FILE, CONFIG_FILE)
            else:
                if os.path.exists(CONFIG_FILE):
                    os.remove(CONFIG_FILE)
            if os.path.exists(CONFIG_BACKUP_FILE):
                os.remove(CONFIG_BACKUP_FILE)
            messagebox.showinfo("Restaurado", "Configurações restauradas. O aplicativo será fechado.")
            self.destroy()

    def set_as_default(self):
        self.save_config_from_ui()
        try:
            with open(DEFAULT_CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(CONFIG, f, ensure_ascii=False, indent=4)
            messagebox.showinfo(
                "Padrão Definido",
                "As configurações atuais foram salvas como o novo padrão do aplicativo.\n\nAo clicar em 'Restaurar Padrões de Fábrica', o app voltará para este estado.",
            )
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível salvar o padrão:\n{e}")

    def open_color_picker(self, var, btn_widget):
        code = colorchooser.askcolor(title="Escolha uma cor", initialcolor=var.get())
        if code and code[1]:
            hex_color = code[1]
            var.set(hex_color)
            r, g, b = code[0]
            text_col = "black" if (r * 299 + g * 587 + b * 114) / 1000 > 125 else "white"
            btn_widget.configure(fg_color=hex_color, text_color=text_col)

    # -------------------------------------------------------------------------
    # Save config from UI
    # -------------------------------------------------------------------------
    def save_config_from_ui(self):
        old_config = dict(CONFIG)

        CONFIG.update({
            "excel_path": self.entryExcel.get(),
            "template_path": self.entryTemplate.get(),
            "save_dir": self.entrySaveDir.get(),
        })

        suggestions_text = self.entrySuggestions.get()
        suggestions_list = [s.strip() for s in suggestions_text.split(",") if s.strip()]
        CONFIG["sugestoes_itens"] = suggestions_list
        self.itemCombobox.configure(values=suggestions_list)

        for _, light_key, dark_key in APPEARANCE_COLOR_FIELDS:
            CONFIG[light_key] = getattr(self, light_key).get()
            CONFIG[dark_key] = getattr(self, dark_key).get()

        for _, key, _, _ in APPEARANCE_NUMBER_FIELDS:
            CONFIG[key] = getattr(self, f"combo_{key}").get()

        CONFIG["fonte_familia"] = self.comboFont.get()
        try:
            CONFIG["fonte_tamanho_base"] = int(self.comboFontSize.get())
        except (TypeError, ValueError):
            CONFIG["fonte_tamanho_base"] = old_config.get("fonte_tamanho_base", 12)
        CONFIG["fonte_negrito_padrao"] = bool(self.switchBold.get())

        accessibility = {
            "tema": self.comboTheme.get(),
            "escala_widget": self.comboWidgetScale.get(),
            "escala_texto": self.comboTextScale.get(),
            "fonte_dislexia": bool(self.switchDyslexia.get()),
            "alto_contraste": bool(self.switchHighContrast.get()),
        }
        CONFIG.update(accessibility)
        save_config(CONFIG)

        visual_changed = any(old_config.get(key) != CONFIG.get(key) for key in RESTART_CONFIG_KEYS)
        if visual_changed:
            if messagebox.askyesno("Reiniciar Agora?", "As alterações visuais requerem reiniciar o aplicativo.\n\nDeseja reiniciar agora?", parent=self):
                with open(CONFIG_BACKUP_FILE, "w", encoding="utf-8") as f:
                    json.dump(old_config, f, ensure_ascii=False, indent=4)
                self.destroy()
                os.execl(sys.executable, sys.executable, *sys.argv)
            else:
                self.statusVar.set("Configurações salvas. Reinicie para aplicar as alterações visuais.")
        else:
            messagebox.showinfo("Sucesso", "Configurações salvas com sucesso.", parent=self)

    # -------------------------------------------------------------------------
    # Generator logic
    # -------------------------------------------------------------------------
    def increase_year(self):
        try:
            year = int(self.selectedYear.get())
            self.selectedYear.set(str(year + 1))
        except ValueError:
            self.selectedYear.set(str(datetime.date.today().year))
        self.on_year_change()

    def decrease_year(self):
        try:
            year = int(self.selectedYear.get())
            self.selectedYear.set(str(year - 1))
        except ValueError:
            self.selectedYear.set(str(datetime.date.today().year))
        self.on_year_change()

    def on_year_change(self, *args):
        self.load_next_number()
        self.statusVar.set("Ano alterado, número atualizado.")

    def load_next_number(self):
        try:
            year_str = self.selectedYear.get()
            if not year_str.isdigit():
                return
            year = int(year_str)
            wb = openpyxl.load_workbook(CONFIG["excel_path"], data_only=True)
            try:
                sheet_name = find_sheet_by_year(wb, year)
                ws = wb[sheet_name]
                self.nextOfficio.set(get_next_number(ws, year))
                self.statusVar.set("Pronto")
            except ValueError:
                # Sheet not found for the given year – assume starting number
                self.nextOfficio.set(f"1/{year}")
                self.statusVar.set("Aba não encontrada. Iniciando em 1.")
        except FileNotFoundError:
            self.nextOfficio.set("ERRO")
            self.statusVar.set("Planilha não encontrada.")
        except Exception:
            self.nextOfficio.set("ERRO")
            self.statusVar.set("Erro ao carregar número.")

    def refresh_item_list(self):
        for widget in self.itemsScroll.winfo_children():
            widget.destroy()

        for idx, item in enumerate(self.items):
            frame = ctk.CTkFrame(self.itemsScroll, fg_color="transparent", border_width=0)
            frame.pack(fill="x", pady=2)

            texto = item["nome"]
            if item["codigo"]:
                texto += f" – {item['codigo']}"
            if item["obs"]:
                texto += f" ({item['obs']})"

            ctk.CTkLabel(frame, text=texto, font=get_font(12)).pack(side="left")
            ctk.CTkButton(
                frame,
                text="X",
                width=30,
                height=24,
                fg_color=COLOR_DANGER,
                text_color=COLOR_CONFIRM_TEXT,
                corner_radius=CORNER_RADIUS,
                border_width=BORDER_WIDTH,
                border_color=COLOR_OUTLINE,
                command=lambda i=idx: self.remove_item_at_index(i),
            ).pack(side="right")

    def add_item(self):
        name = self.itemCombobox.get().strip()
        if not name:
            messagebox.showwarning("Aviso", "O campo 'Item' é obrigatório.")
            return
        code = self.codeEntry.get().strip()
        obs = self.obsEntry.get().strip()
        self.items.append({"nome": name, "codigo": code, "obs": obs})
        self.refresh_item_list()
        self.itemCombobox.set("")
        self.codeEntry.delete(0, tk.END)
        self.obsEntry.delete(0, tk.END)
        self.itemCombobox.focus_set()

    def remove_item_at_index(self, idx):
        if 0 <= idx < len(self.items):
            del self.items[idx]
            self.refresh_item_list()

    def clear_items(self):
        if self.items and messagebox.askyesno("Confirmação", "Deseja realmente limpar toda a lista?"):
            self.items.clear()
            self.refresh_item_list()

    def save_draft(self):
        if not self.items:
            messagebox.showinfo("Rascunho", "Nenhum item adicionado.")
            return
        data = {
            "ano": int(self.selectedYear.get()),
            "oficio_number": self.nextOfficio.get(),
            "itens": self.items,
        }
        try:
            with open(DRAFT_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            self.statusVar.set("Rascunho salvo com sucesso.")
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível salvar o rascunho:\n{str(e)}")

    def load_draft(self):
        if not os.path.exists(DRAFT_FILE):
            messagebox.showinfo("Rascunho", "Nenhum rascunho encontrado.")
            return
        try:
            with open(DRAFT_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            self.selectedYear.set(str(data.get("ano", datetime.date.today().year)))
            self.nextOfficio.set(data.get("oficio_number", ""))
            self.items = data.get("itens", [])
            self.refresh_item_list()
            self.statusVar.set("Rascunho carregado.")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar rascunho:\n{str(e)}")

    def auto_check_draft(self):
        if os.path.exists(DRAFT_FILE):
            if messagebox.askyesno("Rascunho Encontrado", "Existe um rascunho salvo. Deseja carregá-lo?", parent=self):
                self.load_draft()
        self.statusVar.set("Pronto")
        self.focus_force()
        self.itemCombobox.focus_set()

    def delete_draft(self):
        if os.path.exists(DRAFT_FILE):
            if messagebox.askyesno("Excluir", "Deseja excluir o rascunho salvo permanentemente?"):
                try:
                    os.remove(DRAFT_FILE)
                    self.statusVar.set("Rascunho excluído.")
                except Exception as e:
                    messagebox.showerror("Erro", f"Não foi possível excluir o rascunho:\n{e}")
        else:
            messagebox.showinfo("Excluir", "Nenhum rascunho encontrado.")

    def generate_document_action(self):
        if not self.items:
            messagebox.showwarning("Aviso", "Adicione pelo menos um item à lista.")
            return
        number = self.nextOfficio.get()
        if number == "ERRO" or not re.match(r"\d+/\d{4}", number):
            messagebox.showerror("Erro", "Número do ofício inválido.")
            return
        try:
            date_officio = datetime.datetime.strptime(self.dateEntry.get().strip(), "%d/%m/%Y").date()
        except ValueError:
            messagebox.showerror("Erro", "Data inválida. Use dd/mm/aaaa.")
            return
        responsible = self.respEntry.get().strip()
        if not responsible:
            messagebox.showwarning("Aviso", "Preencha o Responsável.")
            return

        if not messagebox.askyesno("Confirmar", f"Gerar ofício {number} com {len(self.items)} itens?"):
            return

        num_part = int(number.split("/")[0])
        year_officio = number.split("/")[-1]
        filename = f"{year_officio} {num_part:03d} - Encaminhamento de Achados e Perdidos.docx"

        default_dir = CONFIG.get("save_dir", SCRIPT_DIR)
        if not os.path.isdir(default_dir):
            default_dir = SCRIPT_DIR

        file_path = filedialog.asksaveasfilename(
            initialdir=default_dir,
            initialfile=filename,
            defaultextension=".docx",
            filetypes=[("Documento Word", "*.docx")],
            title="Salvar Ofício",
        )
        if not file_path:
            return

        self.statusVar.set("Gerando documento...")
        self.update()

        try:
            saved_path = generate_document(number, self.items, date_officio, file_path)
            if not add_excel_row(number, int(year_officio), date_officio, responsible):
                self.statusVar.set("Erro ao salvar planilha")
                return
        except Exception as e:
            messagebox.showerror("Erro", f"Falha na geração:\n{str(e)}")
            self.statusVar.set("Erro")
            return

        self.delete_draft()
        current_num = int(number.split("/")[0])
        self.nextOfficio.set(f"{current_num + 1}/{year_officio}")
        self.items.clear()
        self.refresh_item_list()
        self.itemCombobox.set("")
        self.codeEntry.delete(0, tk.END)
        self.obsEntry.delete(0, tk.END)
        self.respEntry.delete(0, tk.END)
        self.statusVar.set(f"Ofício {number} gerado com sucesso!")
        messagebox.showinfo("Sucesso", f"Ofício salvo em:\n{saved_path}")


if __name__ == "__main__":
    app = OficioGeneratorApp()
    app.mainloop()